from __future__ import annotations

import argparse
import csv
import json
import re
from pathlib import Path

from docx import Document


ELR_CODE_RE = re.compile(r"\[(\d+[a-z]?)\]")


def _clean(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\n", " ").replace("\r", " ")).strip()


def _extract_table_rows(table) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([_clean(cell.text or "") for cell in row.cells])
    return rows


def _extract_elr_changes(table_rows: list[list[str]]) -> list[dict]:
    if not table_rows:
        return []
    header = table_rows[0]
    body = table_rows[1:]
    out: list[dict] = []
    for row in body:
        row = row + [""] * (len(header) - len(row))
        record = {header[i] if i < len(header) else f"col_{i+1}": row[i] for i in range(len(header))}
        elr_name = record.get("ELR名称", "")
        change_desc = record.get("变化说明", "")
        matched_codes = ELR_CODE_RE.findall(elr_name + " " + change_desc)
        record["elr_codes"] = matched_codes
        out.append(record)
    return out


def _extract_element_types(table_rows: list[list[str]]) -> list[dict]:
    if not table_rows:
        return []
    header = table_rows[0]
    body = table_rows[1:]
    out: list[dict] = []
    for row in body:
        row = row + [""] * (len(header) - len(row))
        out.append({header[i] if i < len(header) else f"col_{i+1}": row[i] for i in range(len(header))})
    return out


def _write_tsv(path: Path, rows: list[dict]) -> None:
    if not rows:
        path.write_text("", encoding="utf-8")
        return
    fieldnames = list(rows[0].keys())
    with path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames, delimiter="\t")
        writer.writeheader()
        writer.writerows(rows)


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract structured rules from CAS 2020 background document.")
    parser.add_argument("--input", default="data/taxonomy/cas_2020_elements_background.doc")
    parser.add_argument("--output-json", default="data/taxonomy/cas2020_background_rules.json")
    parser.add_argument("--output-elr-tsv", default="data/taxonomy/cas2020_background_elr_changes.tsv")
    parser.add_argument("--output-types-tsv", default="data/taxonomy/cas2020_background_element_types.tsv")
    args = parser.parse_args()

    doc = Document(args.input)
    if len(doc.tables) < 2:
        raise SystemExit("Expected at least two tables in the background document.")

    elr_rows = _extract_table_rows(doc.tables[0])
    type_rows = _extract_table_rows(doc.tables[1])
    elr_changes = _extract_elr_changes(elr_rows)
    element_types = _extract_element_types(type_rows)

    payload = {
        "source_file": args.input,
        "elr_changes": elr_changes,
        "element_types": element_types,
    }
    Path(args.output_json).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    _write_tsv(Path(args.output_elr_tsv), elr_changes)
    _write_tsv(Path(args.output_types_tsv), element_types)

    print(
        json.dumps(
            {
                "elr_changes": len(elr_changes),
                "element_types": len(element_types),
                "output_json": args.output_json,
                "output_elr_tsv": args.output_elr_tsv,
                "output_types_tsv": args.output_types_tsv,
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
