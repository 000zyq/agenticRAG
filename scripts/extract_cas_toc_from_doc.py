from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document


TOP_RE = re.compile(r"^(CAS\d+)\s+(.+)$")
SUB_RE = re.compile(r"^\[(\d{6})\]\s*(.+)$")


def _clean(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\n", " ").replace("\r", " ")).strip()


def _normalize_sub_name(name: str) -> str:
    text = _clean(name)
    text = re.sub(r"^附注[_:：\-\s]*", "", text)
    text = re.sub(r"（[^）]*）", "", text)
    text = re.sub(r"\([^)]*\)", "", text)
    text = _clean(text)
    return text


def extract_toc(path: Path) -> dict:
    doc = Document(str(path))
    if not doc.tables:
        raise ValueError("No tables found in document.")

    cn_table = doc.tables[0]
    cas_categories: list[dict] = []
    sub_categories: list[dict] = []

    current_cas_code: str | None = None
    current_cas_name: str | None = None

    for row_idx, row in enumerate(cn_table.rows):
        cell0 = _clean(row.cells[0].text or "")
        if not cell0:
            continue

        top_match = TOP_RE.match(cell0)
        if top_match:
            current_cas_code = top_match.group(1)
            current_cas_name = top_match.group(2)
            cas_categories.append(
                {
                    "cas_code": current_cas_code,
                    "cas_name": current_cas_name,
                    "row_index": row_idx,
                    "page_ref": None,
                }
            )
            continue

        sub_match = SUB_RE.match(cell0)
        if sub_match:
            raw_sub_name = sub_match.group(2)
            sub_categories.append(
                {
                    "sub_code": sub_match.group(1),
                    "sub_name_raw": raw_sub_name,
                    "sub_name": _normalize_sub_name(raw_sub_name),
                    "cas_code": current_cas_code,
                    "cas_name": current_cas_name,
                    "row_index": row_idx,
                    "page_ref": None,
                }
            )

    return {
        "source_file": str(path),
        "cas_categories": cas_categories,
        "sub_categories": sub_categories,
    }


def write_tsv(path: Path, cas_categories: list[dict], sub_categories: list[dict]) -> None:
    lines = ["level\tcas_code\tcas_name\tsub_code\tsub_name\tsub_name_raw\trow_index\tpage_ref"]
    for item in cas_categories:
        lines.append(
            "\t".join(
                [
                    "cas",
                    item["cas_code"],
                    item["cas_name"],
                    "",
                    "",
                    "",
                    str(item["row_index"]),
                    "",
                ]
            )
        )
    for item in sub_categories:
        lines.append(
            "\t".join(
                [
                    "sub",
                    item.get("cas_code") or "",
                    item.get("cas_name") or "",
                    item["sub_code"],
                    item["sub_name"],
                    item.get("sub_name_raw") or "",
                    str(item["row_index"]),
                    "",
                ]
            )
        )
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract CAS TOC-like hierarchy from Word element list.")
    parser.add_argument("--input", required=True, help="Path to .doc/.docx file.")
    parser.add_argument("--output-json", default="data/taxonomy/cas2020_toc.json")
    parser.add_argument("--output-tsv", default="data/taxonomy/cas2020_toc.tsv")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        raise SystemExit(f"Input not found: {input_path}")

    payload = extract_toc(input_path)
    output_json = Path(args.output_json)
    output_tsv = Path(args.output_tsv)
    output_json.parent.mkdir(parents=True, exist_ok=True)

    output_json.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    write_tsv(output_tsv, payload["cas_categories"], payload["sub_categories"])

    print(
        json.dumps(
            {
                "cas_categories": len(payload["cas_categories"]),
                "sub_categories": len(payload["sub_categories"]),
                "output_json": str(output_json),
                "output_tsv": str(output_tsv),
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
